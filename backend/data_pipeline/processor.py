#!/usr/bin/env python3
"""
Document processor module for extracting and preprocessing text from documents.

This module handles:
1. Text extraction from various file formats
2. Basic preprocessing and normalization
3. Document chunking for long documents
"""

import os
import logging
from pathlib import Path
from typing import Dict, Any, Tuple, Union, List, Callable

logger = logging.getLogger("DocumentProcessor")

class DocumentProcessor:
    """Extract and preprocess text from document files"""
    
    def __init__(self):
        self.supported_extensions = {
            ".txt": self._extract_from_txt,
            ".pdf": self._extract_from_pdf,
            ".docx": self._extract_from_docx,
            ".json": self._extract_from_json,
            ".html": self._extract_from_html
        }
        self.max_length = 8000  # Maximum text length to process
    
    def process(self, file_path: Union[str, Path]) -> Tuple[str, Dict[str, Any]]:
        """
        Extract and preprocess text from a file
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (preprocessed_text, metadata)
        """
        text, metadata = self._extract(file_path)
        if text:
            text = self._preprocess(text)
        return text, metadata
    
    def _extract(self, file_path: Union[str, Path]) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from a file
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (text, metadata)
        """
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_extensions:
            logger.warning(f"Unsupported file extension: {file_extension}")
            return "", {}
        
        try:
            extractor = self.supported_extensions[file_extension]
            text, metadata = extractor(file_path)
            
            # Add common metadata
            metadata.update({
                "filename": file_path.name,
                "file_path": str(file_path),
                "file_extension": file_extension,
                "file_size_bytes": file_path.stat().st_size,
            })
            
            return text, metadata
        except Exception as e:
            logger.error(f"Error extracting from {file_path}: {e}")
            return "", {}
    
    def _preprocess(self, text: str) -> str:
        """
        Preprocess text for embedding
        
        Args:
            text: The input text
            
        Returns:
            Preprocessed text
        """
        if not text:
            return ""
        
        # Basic cleaning
        text = text.strip()
        
        # Replace multiple newlines with single newline
        text = '\n'.join([line.strip() for line in text.split('\n') if line.strip()])
        
        # Replace multiple spaces with single space
        text = ' '.join(text.split())
        
        # Truncate if too long
        if len(text) > self.max_length:
            logger.warning(f"Text truncated from {len(text)} to {self.max_length} characters")
            text = text[:self.max_length]
        
        return text
    
    def split_into_chunks(self, text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """
        Split text into overlapping chunks for better processing
        
        Args:
            text: The input text
            chunk_size: Maximum size of each chunk
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks
        """
        if not text or len(text) <= chunk_size:
            return [text] if text else []
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Find the end of the chunk
            end = start + chunk_size
            
            # Adjust end to avoid cutting in the middle of a word or sentence
            if end < len(text):
                # Try to find a sentence boundary
                sentence_end = text.rfind('. ', start, end)
                if sentence_end > start + chunk_size // 2:
                    end = sentence_end + 2  # Include the period and space
                else:
                    # Try to find a word boundary
                    word_end = text.rfind(' ', start, end)
                    if word_end > start + chunk_size // 2:
                        end = word_end + 1  # Include the space
            
            chunks.append(text[start:end])
            start = end - overlap
        
        return chunks
    
    def _extract_from_txt(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a plain text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Basic metadata extraction for text files
        metadata = {
            "content_type": "text/plain"
        }
        
        # Try to extract title and other metadata from legal documents
        lines = text.split('\n')
        for i, line in enumerate(lines[:10]):  # Check first 10 lines
            line = line.strip()
            if line.startswith("CASE:") or line.startswith("TITLE:"):
                metadata["title"] = line.split(":", 1)[1].strip()
            elif line.startswith("CITATION:"):
                metadata["citation"] = line.split(":", 1)[1].strip()
            elif line.startswith("SECTION:"):
                metadata["section"] = line.split(":", 1)[1].strip()
        
        return text, metadata
    
    def _extract_from_pdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a PDF file"""
        # Use PyPDF2 or pdfminer.six if available
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                
                metadata = {
                    "content_type": "application/pdf",
                    "page_count": len(reader.pages)
                }
                
                # Try to extract PDF metadata
                if reader.metadata:
                    for key, value in reader.metadata.items():
                        if key and value and isinstance(value, str):
                            metadata[key.lower().replace('/', '_')] = value
                
                return text, metadata
        except ImportError:
            logger.warning("PyPDF2 not available. Install with: pip install PyPDF2")
            return f"[PDF content from {file_path}]", {"content_type": "application/pdf"}
    
    def _extract_from_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a DOCX file"""
        try:
            import docx
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            metadata = {
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "paragraph_count": len(doc.paragraphs)
            }
            
            # Try to extract document properties
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                if hasattr(props, 'title') and props.title:
                    metadata["title"] = props.title
                if hasattr(props, 'author') and props.author:
                    metadata["author"] = props.author
                if hasattr(props, 'created') and props.created:
                    metadata["created"] = props.created.isoformat()
            
            return text, metadata
        except ImportError:
            logger.warning("python-docx not available. Install with: pip install python-docx")
            return f"[DOCX content from {file_path}]", {"content_type": "application/docx"}
    
    def _extract_from_json(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a JSON file"""
        import json
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Extract text and metadata based on common legal document formats
        text = ""
        metadata = {"content_type": "application/json"}
        
        # Handle different JSON structures
        if isinstance(data, dict):
            # Extract content from known fields
            for content_field in ["content", "text", "body", "document"]:
                if content_field in data and isinstance(data[content_field], str):
                    text = data[content_field]
                    break
            
            # Extract metadata from known fields
            for meta_field in ["title", "citation", "section", "id", "author", "date"]:
                if meta_field in data and isinstance(data[meta_field], (str, int, float)):
                    metadata[meta_field] = data[meta_field]
        
        return text, metadata
    
    def _extract_from_html(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from an HTML file"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            
            # Extract text
            text = soup.get_text(separator='\n')
            
            # Extract metadata
            metadata = {"content_type": "text/html"}
            
            # Try to get title
            if soup.title and soup.title.string:
                metadata["title"] = soup.title.string.strip()
            
            # Extract meta tags
            for meta in soup.find_all('meta'):
                if meta.get('name') and meta.get('content'):
                    metadata[meta['name'].lower()] = meta['content']
            
            return text, metadata
        except ImportError:
            logger.warning("BeautifulSoup not available. Install with: pip install beautifulsoup4")
            return f"[HTML content from {file_path}]", {"content_type": "text/html"} 